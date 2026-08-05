"""Generates specwrite/templates/blank_spec_template.docx — the structural
placeholder used by create_blank_spec() for "New Spec -> from template".

This is a *synthetic* placeholder: correct section structure, headings, and
table shapes, but no real branding/boilerplate. It contains no customer
data, so unlike real spec documents it's safe to commit to the repo. Swap
in a real blank Toppan master template by regenerating this file (or just
replacing it directly) once one is available — nothing else needs to
change, since create_blank_spec() only depends on the section structure
being parseable.

Run with: python scripts/build_blank_template.py
"""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.shared import Inches, RGBColor

RED = RGBColor(0xFF, 0x00, 0x00)
OUTPUT_PATH = Path(__file__).resolve().parent.parent / "specwrite" / "templates" / "blank_spec_template.docx"


def _heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.color.rgb = RED
    run.bold = True


def _fill(table, rows: list[list[str]]) -> None:
    for r, row_values in enumerate(rows):
        for c, value in enumerate(row_values):
            table.cell(r, c).text = value


def build() -> None:
    doc = Document()

    header = doc.sections[0].header
    pd = header.add_table(rows=4, cols=4, width=Inches(6))
    _fill(
        pd,
        [
            ["Customer:", "", "Spec #:", ""],
            ["Item:", "", "Date of Issue:", ""],
            ["", "", "Revision #:", "01"],
            ["", "", "Structure Code:", ""],
        ],
    )
    pd.cell(2, 0).merge(pd.cell(3, 0))
    pd.cell(2, 1).merge(pd.cell(3, 1))
    pd.cell(2, 0).text = "Structure Description:"
    pd.cell(2, 1).text = ""

    _heading(doc, "Locations")
    t = doc.add_table(rows=1, cols=4)
    _fill(t, [["Customer Location:", "", "Facility:", ""]])

    _heading(doc, "Bill of Materials")
    doc.add_table(rows=1, cols=6).rows[0].cells[0].text = "Basis Wt\n(#/ream)"
    bom = doc.tables[-1]
    for i, h in enumerate(["Basis Wt\n(#/ream)", "Basis Wt range", "Caliper\n(mils)", "Raw Material", "Supplier", "Raw Material Item Code"]):
        bom.cell(0, i).text = h

    _heading(doc, "Secondary Approved Materials")
    sam = doc.add_table(rows=1, cols=6)
    for i, h in enumerate(["Basis Wt\n(#/ream)", "Basis Wt range", "Caliper\n(mils)", "Raw Material", "Supplier", "Raw Material Item Code"]):
        sam.cell(0, i).text = h

    _heading(doc, "Process Routing")
    pr = doc.add_table(rows=1, cols=5)
    for i, h in enumerate(["Pass", "Process", "Machine", "Process Conditions", "Comments"]):
        pr.cell(0, i).text = h

    _heading(doc, "Physical Attributes & Testing")
    pat = doc.add_table(rows=1, cols=10)
    for i, h in enumerate(["Attribute", "COA", "Units", "LSL", "LCL", "Target", "UCL", "USL", "Procedure", "Frequency"]):
        pat.cell(0, i).text = h

    _heading(doc, "Slitting Information")
    t = doc.add_table(rows=1, cols=6)
    _fill(t, [["Slit Width:", "", "Core ID:", "", "Splice Tape:", ""]])

    _heading(doc, "Packing Information")
    t = doc.add_table(rows=1, cols=6)
    _fill(t, [["Pallet Size:", "", "Poly Bags:", "", "Pallet Labels:", ""]])

    _heading(doc, "Reporting Requirements")
    t = doc.add_table(rows=1, cols=4)
    _fill(t, [["COA Required?:", "", "COA Recipient(s):", ""]])

    _heading(doc, "Customer Sampling Requirements")
    t = doc.add_table(rows=1, cols=4)
    _fill(t, [["Samples Required?:", "", "Send to:", ""]])

    _heading(doc, "Revision History")
    rev = doc.add_table(rows=2, cols=4)
    _fill(
        rev,
        [
            ["Revision #", "Who", "Date", "Revision"],
            ["01", "", "", "Spec created from template."],
        ],
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
