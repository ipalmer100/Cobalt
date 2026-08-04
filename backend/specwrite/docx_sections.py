"""Read-side parser: locate each spec's 11 sections inside a .docx and
flatten them into structured tables.

Detection strategy (generalized from the org's existing BOM-extraction VBA):
for each known section name, find the paragraph whose text matches it, then
take the *next table that follows it* in document order. Product Description
is the exception — it lives in the page header as a label/value grid, with a
fallback to the first body table for older specs that never got a header
(mirrors the VBA's ``headerNum = 0`` branch).
"""

from __future__ import annotations

import re

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from .models import ParsedTable, Spec, TableShape

# Canonical section names, in the order they appear in a spec.
PRODUCT_DESCRIPTION = "Product Description"

BODY_SECTIONS = [
    "Locations",
    "Bill of Materials",
    "Secondary Approved Materials",
    "Process Routing",
    "Physical Attributes & Testing",
    "Slitting Information",
    "Packing Information",
    "Reporting Requirements",
    "Customer Sampling Requirements",
    "Revision History",
]

ALL_SECTIONS = [PRODUCT_DESCRIPTION, *BODY_SECTIONS]

_FIELD_GRID_MIN_COLON_FRACTION = 0.25


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip().lower()


_NORMALIZED_BODY_SECTIONS = {_normalize(name): name for name in BODY_SECTIONS}


def _table_to_rows(table: Table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    return rows


def _classify_shape(rows: list[list[str]]) -> TableShape:
    cells = [c for row in rows for c in row if c]
    if not cells:
        return TableShape.RECORDS
    colon_count = sum(1 for c in cells if c.rstrip().endswith(":"))
    fraction = colon_count / len(cells)
    return TableShape.FIELDS if fraction >= _FIELD_GRID_MIN_COLON_FRACTION else TableShape.RECORDS


def _build_parsed_table(section: str, table: Table, table_index: int, location: str) -> ParsedTable:
    rows = _table_to_rows(table)
    shape = _classify_shape(rows)
    header_row = rows[0] if shape == TableShape.RECORDS and rows else None
    return ParsedTable(
        section=section,
        table_index=table_index,
        location=location,
        shape=shape,
        header_row=header_row,
        rows=rows,
    )


def _iter_body_blocks(doc: Document):
    """Yield ('p', Paragraph) / ('tbl', Table, table_index) in document order."""
    body = doc.element.body
    paragraphs = doc.paragraphs
    tables = doc.tables
    p_idx = 0
    t_idx = 0
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            yield ("p", paragraphs[p_idx])
            p_idx += 1
        elif tag == "tbl":
            yield ("tbl", tables[t_idx], t_idx)
            t_idx += 1


def find_body_section_tables(doc: Document) -> dict[str, ParsedTable]:
    """Pair each known section title paragraph with the table that follows it."""
    found: dict[str, ParsedTable] = {}
    pending_section: str | None = None

    for block in _iter_body_blocks(doc):
        if block[0] == "p":
            paragraph: Paragraph = block[1]
            text = _normalize(paragraph.text)
            if text in _NORMALIZED_BODY_SECTIONS:
                pending_section = _NORMALIZED_BODY_SECTIONS[text]
        else:
            _, table, table_index = block
            if pending_section and pending_section not in found:
                found[pending_section] = _build_parsed_table(pending_section, table, table_index, "body")
            pending_section = None

    return found


def extract_product_description(doc: Document) -> ParsedTable:
    """Product Description normally lives in the page header's first table;
    fall back to the first body table when a spec has no header table at all
    (older specs, per the VBA's ``headerNum = 0`` branch)."""
    header = doc.sections[0].header
    if header.tables:
        return _build_parsed_table(PRODUCT_DESCRIPTION, header.tables[0], 0, "header")
    if doc.tables:
        return _build_parsed_table(PRODUCT_DESCRIPTION, doc.tables[0], 0, "body")
    return ParsedTable(
        section=PRODUCT_DESCRIPTION,
        table_index=-1,
        location="missing",
        shape=TableShape.FIELDS,
        header_row=None,
        rows=[],
    )


def parse_document(path: str) -> Spec:
    doc = Document(path)

    tables: dict[str, ParsedTable] = {}
    warnings: list[str] = []

    tables[PRODUCT_DESCRIPTION] = extract_product_description(doc)
    tables.update(find_body_section_tables(doc))

    for name in ALL_SECTIONS:
        if name not in tables:
            warnings.append(f"Section not found: {name}")

    pd_fields = tables[PRODUCT_DESCRIPTION].fields()
    spec_number = pd_fields.get("Spec #", "")
    customer = pd_fields.get("Customer", "")
    revision_number = pd_fields.get("Revision #", "")

    return Spec(
        file_path=path,
        spec_number=spec_number,
        customer=customer,
        revision_number=revision_number,
        tables=tables,
        warnings=warnings,
    )
