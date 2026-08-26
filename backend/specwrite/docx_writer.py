"""Write-side: address-based cell writes, generalized from wrdRevision.bas.

The existing VBA tools write two ways: most of them do a blind Word
``Find/Replace`` scoped to a table's range (fragile — breaks if the same
text appears twice, and can't target "this specific cell" independent of
its current value). Only the revision-history macro writes by cell address
(``tbl.Cell(r, c).Range.Text = value``). This module generalizes that
address-based approach to every table, and preserves the target cell's
existing run formatting (font, color, bold) instead of resetting it to a
default run — a docx analog of editing one field of a markdown frontmatter
block without touching the rest of the file.
"""

from __future__ import annotations

import io
import os
import re
import threading
from contextlib import contextmanager
from datetime import date as date_cls
from pathlib import Path

from docx import Document
from docx.table import Table, _Cell

from .docx_sections import PRODUCT_DESCRIPTION, extract_product_description, find_body_section_tables
from .models import ParsedTable

# One lock per spec file. Every write here is read-modify-write over the
# whole .docx, so two overlapping writes to one file both read the old
# bytes -- and, worse, both stream a fresh zip into the same path at once,
# interleaving into a file Word can no longer open. That is reachable in
# ordinary use: commit a cell, click straight into the next one and commit
# that too, and the second save starts before the first has finished.
_file_locks: dict[str, threading.Lock] = {}
_file_locks_guard = threading.Lock()


def _lock_key(path: str) -> str:
    """Same file, same lock, however the path was spelled."""
    try:
        return os.path.normcase(str(Path(path).resolve()))
    except OSError:
        return os.path.normcase(os.path.abspath(path))


@contextmanager
def _exclusive(path: str):
    key = _lock_key(path)
    with _file_locks_guard:
        lock = _file_locks.setdefault(key, threading.Lock())
    with lock:
        yield


def _save_atomically(doc, path: str) -> None:  # noqa: ANN001
    """Write to a sibling temp file, then rename over the original.

    ``doc.save(path)`` truncates the customer's document and then streams a
    new zip into it, so anything that interrupts it -- a crash, the app
    being closed, a sync client reading mid-write -- leaves a spec that no
    longer opens, with no copy of what was there before. ``os.replace`` is
    atomic on both POSIX and Windows: the file is either wholly the old
    version or wholly the new one.
    """
    target = Path(path)
    temp = target.with_name(f".{target.name}.specwrite-tmp")
    try:
        doc.save(str(temp))
        os.replace(temp, target)
    except BaseException:
        temp.unlink(missing_ok=True)
        raise


def set_cell_text(cell: _Cell, value: str) -> None:
    """Set a cell's text while preserving the formatting of its first run."""
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.text = value
        return

    first = paragraphs[0]
    if first.runs:
        first.runs[0].text = value
        for extra_run in first.runs[1:]:
            extra_run._element.getparent().remove(extra_run._element)
    else:
        first.text = value

    for extra_para in paragraphs[1:]:
        extra_para._element.getparent().remove(extra_para._element)


def write_record_cell(table: Table, row: int, col: int, value: str) -> None:
    """Write a single cell by (row, col) address — the mass-edit grid's
    primary write primitive. Merge-safe: python-docx's ``table.cell()``
    already resolves a merged span's grid position to its origin cell."""
    set_cell_text(table.cell(row, col), value)


def write_field(table: Table, label: str, value: str) -> bool:
    """Find a ``Label:`` cell (case-insensitive, colon optional in the
    argument) in a FIELDS-shape table and set the value cell that follows
    it in the same row. Returns False if the label wasn't found."""
    target = label.strip().rstrip(":").strip().lower()
    for row in table.rows:
        cells = row.cells
        for i, cell in enumerate(cells):
            text = cell.text.strip()
            if text.rstrip(":").strip().lower() == target and text.endswith(":"):
                if i + 1 < len(cells):
                    set_cell_text(cells[i + 1], value)
                    return True
    return False


def add_record_row(table: Table, values: list[str]) -> None:
    """Append a data row to a RECORDS-shape table, copying run formatting
    from the previous last row so the new row doesn't fall back to Word's
    bare default style."""
    template_row = table.rows[-1] if len(table.rows) > 1 else None
    new_row = table.add_row()
    for i, cell in enumerate(new_row.cells):
        if i < len(values):
            if template_row is not None and i < len(template_row.cells):
                _copy_run_format(template_row.cells[i], cell)
            set_cell_text(cell, values[i])


def _copy_run_format(source_cell: _Cell, target_cell: _Cell) -> None:
    source_runs = source_cell.paragraphs[0].runs if source_cell.paragraphs else []
    if not source_runs or not target_cell.paragraphs:
        return
    source_rpr = source_runs[0]._element.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
    if source_rpr is None:
        return
    target_para = target_cell.paragraphs[0]
    if not target_para.runs:
        target_para.add_run("")
    target_run_element = target_para.runs[0]._element
    existing_rpr = target_run_element.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
    if existing_rpr is not None:
        target_run_element.remove(existing_rpr)
    import copy

    target_run_element.insert(0, copy.deepcopy(source_rpr))


_TRAILING_DIGITS = re.compile(r"(\d+)$")


def _next_revision_number(previous: str) -> str:
    """Increment a revision number string, preserving its zero-padded
    width (e.g. "06" -> "07", "9" -> "10", "" -> "01")."""
    match = _TRAILING_DIGITS.search(previous.strip())
    if not match:
        return "01"
    digits = match.group(1)
    incremented = str(int(digits) + 1)
    if len(incremented) < len(digits):
        incremented = incremented.zfill(len(digits))
    return incremented


def _resolve_table(doc: Document, section: str, table_index: int | None = None) -> Table:
    """Locate a section's table in an already-open Document, the same way
    the parser does, so writes target exactly what the reader last showed.

    ``table_index`` disambiguates when a spec holds more than one table for
    the section (a Duplex and a Triplex Process Routing). Section name alone
    would silently write every such row into whichever table came first, so
    the mass-edit grid always carries the index the row was read from.
    """
    if section == PRODUCT_DESCRIPTION:
        parsed = extract_product_description(doc)
        if parsed.location == "missing":
            raise ValueError(f"{PRODUCT_DESCRIPTION} table not found")
        return doc.sections[0].header.tables[0] if parsed.location == "header" else doc.tables[parsed.table_index]

    if table_index is not None and table_index >= 0:
        if table_index >= len(doc.tables):
            raise ValueError(f"Table index {table_index} out of range for {section}")
        return doc.tables[table_index]

    body_tables, _ = find_body_section_tables(doc)
    parsed = body_tables.get(section)
    if not parsed:
        raise ValueError(f"Section not found: {section}")
    return doc.tables[parsed[0].table_index]


def apply_to_bytes(data: bytes, mutate) -> bytes:  # noqa: ANN001
    """Open a spec held in memory, let `mutate(doc)` change it, hand back the
    new bytes.

    The storage-neutral write primitive: a SharePoint document is downloaded,
    edited and uploaded, never touching a filesystem. Every path-based
    function below is this plus a read and a write.
    """
    doc = Document(io.BytesIO(data))
    mutate(doc)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def write_cell(path: str, section: str, row: int, col: int, value: str, table_index: int | None = None) -> None:
    """Open a spec, write one cell by address in the named section, save.
    The mass-edit grid's primary write path."""
    with _exclusive(path):
        doc = Document(path)
        table = _resolve_table(doc, section, table_index)
        write_record_cell(table, row, col, value)
        _save_atomically(doc, path)


def write_field_value(path: str, section: str, label: str, value: str, table_index: int | None = None) -> bool:
    """Open a spec, write one ``Label:`` field's value in the named section, save."""
    with _exclusive(path):
        doc = Document(path)
        table = _resolve_table(doc, section, table_index)
        found = write_field(table, label, value)
        if found:
            _save_atomically(doc, path)
        return found


def append_row(path: str, section: str, values: list[str], table_index: int | None = None) -> None:
    """Open a spec, append a data row to a RECORDS-shape section, save."""
    with _exclusive(path):
        doc = Document(path)
        table = _resolve_table(doc, section, table_index)
        add_record_row(table, values)
        _save_atomically(doc, path)


def write_edits_batch(edits: list[dict]) -> None:
    """Apply many cell/field writes in one pass, grouped by file so each
    .docx is opened and saved exactly once no matter how many cells in it
    changed — the primitive the fill-handle drag needs to feel instant
    even when it spans hundreds of rows across dozens of files, instead
    of paying a full parse+serialize round trip per cell.

    Each edit dict: {"path", "section", "kind": "record"|"field",
    "row", "col", "label", "value", "table_index"}.
    """
    by_path: dict[str, list[dict]] = {}
    for edit in edits:
        by_path.setdefault(edit["path"], []).append(edit)

    for path, path_edits in by_path.items():
        with _exclusive(path):
            doc = Document(path)
            table_cache: dict[tuple[str, int | None], Table] = {}
            for edit in path_edits:
                section = edit["section"]
                table_index = edit.get("table_index")
                cache_key = (section, table_index)
                table = table_cache.get(cache_key)
                if table is None:
                    table = _resolve_table(doc, section, table_index)
                    table_cache[cache_key] = table
                if edit["kind"] == "record":
                    write_record_cell(table, edit["row"], edit["col"], edit["value"])
                else:
                    write_field(table, edit["label"], edit["value"])
            _save_atomically(doc, path)


def clear_records(path: str, section: str, table_index: int | None = None) -> None:
    """Remove every data row from a RECORDS-shape section, keeping only
    the header row. Used when spinning up a new spec (duplicate or blank
    template) so its Revision History doesn't inherit another spec's log."""
    with _exclusive(path):
        doc = Document(path)
        table = _resolve_table(doc, section, table_index)
        for row in list(table.rows[1:]):
            row._tr.getparent().remove(row._tr)
        _save_atomically(doc, path)


def apply_revision(path: str, who: str, revision_text: str, revision_date: date_cls | None = None) -> str:
    """Append a Revision History row and bump the Revision # in Product
    Description. Returns the new revision number. Generalizes
    wrdRevision.bas's ``tbl.Rows.Add`` + header cell-address write."""
    # Held for the whole read-modify-write: the new revision number is
    # derived from the last row that's currently there, so two revisions
    # racing would both read the same previous number and one would be
    # lost -- taking the Revision # / Revision History pairing with it.
    with _exclusive(path):
        doc = Document(path)
        body_tables, _ = find_body_section_tables(doc)
        rev_tables = body_tables.get("Revision History")
        if not rev_tables:
            raise ValueError("Revision History section not found")

        rev_table = doc.tables[rev_tables[0].table_index]
        last_row_values = [c.text.strip() for c in rev_table.rows[-1].cells]
        previous_rev = last_row_values[0] if last_row_values else ""
        new_rev = _next_revision_number(previous_rev)
        rev_date = (revision_date or date_cls.today()).strftime("%m/%d/%Y")

        add_record_row(rev_table, [new_rev, who, rev_date, revision_text])

        pd_parsed: ParsedTable = extract_product_description(doc)
        if pd_parsed.location != "missing":
            source = doc.sections[0].header.tables[0] if pd_parsed.location == "header" else doc.tables[pd_parsed.table_index]
            write_field(source, "Revision #", new_rev)

        _save_atomically(doc, path)
        return new_rev
