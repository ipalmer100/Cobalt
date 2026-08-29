"""The structure-only export.

The point of this tool is what it leaves out. The first test is therefore
the one that matters: take specs stuffed with distinctive values and prove
none of them -- nor the customer names in the folder tree, nor the file
names -- reach the report.
"""

import json

from docx import Document

from cobalt.structure_export import export_folder, main

from .fixtures.builder import build_sample_spec_docx

# Values planted in the documents. Anything appearing in the report is a
# leak; they are nonsense strings so a match cannot be a coincidence.
SECRETS = ["ZZQWERTYVALUE", "PLANTCODE99887", "PRICEPERTHOUSAND4213"]


def _spec_with_secrets(path, spec_number="SW0001"):
    build_sample_spec_docx(str(path), spec_number=spec_number)
    doc = Document(str(path))
    # tables[0] is Locations; the Bill of Materials is the one after it.
    bom = doc.tables[1]
    bom.rows[1].cells[1].text = SECRETS[0]
    bom.rows[1].cells[2].text = SECRETS[1]
    for row in doc.tables[2].rows[1:]:
        row.cells[0].text = SECRETS[2]
    doc.save(str(path))


def test_no_cell_value_reaches_the_report(tmp_path):
    vault = tmp_path / "Specs"
    (vault / "Very Secret Customer Ltd").mkdir(parents=True)
    _spec_with_secrets(vault / "Very Secret Customer Ltd" / "CONFIDENTIAL-EG0614.docx")

    report, _ = export_folder(str(vault))
    text = json.dumps(report.to_json())

    for secret in SECRETS:
        assert secret not in text
    # Nor the folder or file names, which carry the customer.
    assert "Very Secret Customer" not in text
    assert "CONFIDENTIAL" not in text
    assert report.to_json()["contains_cell_values"] is False


def test_the_root_is_named_but_nothing_below_it(tmp_path):
    """The vault path is the user's own folder and tells them which run this
    was; the tree beneath it is what carries customer names."""
    vault = tmp_path / "Specs"
    (vault / "Acme Foods" / "48g PET").mkdir(parents=True)
    build_sample_spec_docx(str(vault / "Acme Foods" / "48g PET" / "EG1.docx"))

    data, _ = export_folder(str(vault))
    text = json.dumps(data.to_json())
    assert "Acme Foods" not in text
    assert "48g PET" not in text


def test_structure_is_actually_described(tmp_path):
    vault = tmp_path / "Specs"
    vault.mkdir()
    build_sample_spec_docx(str(vault / "a.docx"))

    report, _ = export_folder(str(vault))
    data = report.to_json()

    assert data["vault"]["specs_analyzed"] == 1
    assert data["sections_seen"]["Bill of Materials"] == 1
    layout = data["layouts"][0]
    bom = next(t for t in layout["tables"] if t["section"] == "Bill of Materials")
    assert bom["shape"] == "records"
    assert "Raw Material" in [c["label"] for c in bom["columns"]]
    assert bom["column_count"] == 5
    # Fill rates are numbers about the values, not the values.
    assert all(0 <= c["filled_pct"] <= 100 for c in bom["columns"])


def test_specs_sharing_a_layout_are_grouped(tmp_path):
    vault = tmp_path / "Specs"
    vault.mkdir()
    for i in range(4):
        build_sample_spec_docx(str(vault / f"spec{i}.docx"), spec_number=f"SW000{i}")

    report, _ = export_folder(str(vault))
    data = report.to_json()
    assert data["vault"]["distinct_layouts"] == 1
    assert data["layouts"][0]["specs"] == 4
    assert len(data["layouts"][0]["example_ids"]) == 4


def test_a_differing_layout_lands_in_its_own_group(tmp_path):
    vault = tmp_path / "Specs"
    vault.mkdir()
    build_sample_spec_docx(str(vault / "normal.docx"))
    odd = vault / "odd.docx"
    build_sample_spec_docx(str(odd))
    doc = Document(str(odd))
    doc.tables[1].rows[0].cells[3].text = "Vendor Code"  # renamed BOM column
    doc.save(str(odd))

    report, _ = export_folder(str(vault))
    data = report.to_json()
    assert data["vault"]["distinct_layouts"] == 2
    labels = {
        c["label"]
        for layout in data["layouts"]
        for table in layout["tables"]
        for c in table["columns"]
    }
    assert {"Designation", "Vendor Code"} <= labels


def test_a_misdetected_header_is_flagged_rather_than_hidden(tmp_path):
    """The case worth surfacing: a header cell left blank, or the same text
    repeated across columns, is how a data row ends up read as a header."""
    vault = tmp_path / "Specs"
    vault.mkdir()
    path = vault / "odd.docx"
    build_sample_spec_docx(str(path))
    doc = Document(str(path))
    doc.tables[1].rows[0].cells[0].text = ""
    doc.tables[1].rows[0].cells[3].text = "Supplier"  # duplicate of column 2
    doc.save(str(path))

    data, _ = export_folder(str(vault))
    suspects = data.to_json()["suspect_labels"]
    assert "Column 1" in suspects
    assert "Supplier (2)" in suspects


def test_inactive_specs_are_counted_without_naming_the_folder(tmp_path):
    vault = tmp_path / "Specs"
    (vault / "Inactive Specifications" / "Acme").mkdir(parents=True)
    (vault / "Acme").mkdir(parents=True)
    build_sample_spec_docx(str(vault / "Acme" / "live.docx"))
    build_sample_spec_docx(str(vault / "Inactive Specifications" / "Acme" / "retired.docx"))

    report, _ = export_folder(str(vault))
    data = report.to_json()
    assert data["vault"]["specs_analyzed"] == 2
    assert data["vault"]["filed_as_inactive"] == 1


def test_an_unreadable_file_is_reported_and_the_walk_continues(tmp_path):
    vault = tmp_path / "Specs"
    vault.mkdir()
    build_sample_spec_docx(str(vault / "good.docx"))
    (vault / "broken.docx").write_bytes(b"not a docx at all")

    report, _ = export_folder(str(vault))
    data = report.to_json()
    assert data["vault"]["specs_analyzed"] == 1
    assert data["vault"]["unreadable"] == 1
    # The failure names the spec by id, not by path.
    assert "broken" not in json.dumps(data["unreadable_specs"])


def test_word_lock_files_and_in_flight_saves_are_skipped(tmp_path):
    vault = tmp_path / "Specs"
    vault.mkdir()
    build_sample_spec_docx(str(vault / "real.docx"))
    (vault / "~$real.docx").write_bytes(b"lock")
    (vault / ".real.docx.cobalt-tmp").write_bytes(b"partial")

    report, _ = export_folder(str(vault))
    assert report.to_json()["vault"]["specs_analyzed"] == 1


def test_ids_are_stable_between_runs(tmp_path):
    vault = tmp_path / "Specs"
    vault.mkdir()
    build_sample_spec_docx(str(vault / "a.docx"))

    first, map_a = export_folder(str(vault))
    second, map_b = export_folder(str(vault))
    assert map_a == map_b
    assert first.to_json()["layouts"][0]["example_ids"] == second.to_json()["layouts"][0]["example_ids"]


def test_limit_stops_early(tmp_path):
    vault = tmp_path / "Specs"
    vault.mkdir()
    for i in range(5):
        build_sample_spec_docx(str(vault / f"s{i}.docx"))

    report, _ = export_folder(str(vault), limit=2)
    assert report.to_json()["vault"]["specs_analyzed"] == 2


def test_cli_writes_the_report_and_a_local_only_mapping(tmp_path, capsys, monkeypatch):
    vault = tmp_path / "Specs"
    (vault / "Acme Foods").mkdir(parents=True)
    _spec_with_secrets(vault / "Acme Foods" / "EG0614.docx")

    monkeypatch.chdir(tmp_path)
    out = tmp_path / "report.json"
    assert main([str(vault), "--out", str(out)]) == 0

    data = json.loads(out.read_text())
    assert data["cobalt_structure_report"] == 1
    assert data["vault"]["specs_analyzed"] == 1
    for secret in SECRETS:
        assert secret not in out.read_text()

    mapping = json.loads((tmp_path / "report-local-map.json").read_text())
    # The mapping is the one place a real path appears -- that is its job,
    # and it is why it is a separate file.
    assert list(mapping.values()) == ["Acme Foods/EG0614.docx"] or list(mapping.values()) == [
        "Acme Foods\\EG0614.docx"
    ]

    printed = capsys.readouterr().out
    assert "KEEP THIS ONE LOCAL" in printed


def test_cli_rejects_an_unknown_flag(tmp_path):
    assert main([str(tmp_path), "--send-everything"]) == 2
