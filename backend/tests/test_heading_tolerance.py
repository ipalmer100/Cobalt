"""Headings the archive actually contains, and what they must classify as.

Every string here was taken from a structure report over a real 1,811-spec
library. The ones that used to fail were not exotic -- a stray character
typed into the heading paragraph, a dropped letter, a bilingual plant --
but each one made a whole section of that spec invisible in the app.

The refusals matter as much as the matches: a heading guessed wrong routes
writes into the wrong table, which is worse than one sent to the exception
queue for a human to place.
"""

import pytest

from cobalt.docx_sections import BODY_SECTIONS, classify_heading

# (heading as written, section it must resolve to)
RECOVERED = [
    # Stray characters glued to the front.
    ("`Locations", "Locations"),
    ("eLocations", "Locations"),
    ("KLocations", "Locations"),
    ("mLocations", "Locations"),
    ("L75Locations", "Locations"),
    ("2.99Locations", "Locations"),
    ("snippinLocations", "Locations"),
    ("-Locations", "Locations"),
    (".Bill of Materials", "Bill of Materials"),
    ("oProcess Routing", "Process Routing"),
    ("9Revision History", "Revision History"),
    # Trailing punctuation.
    ("Bill of Materials-", "Bill of Materials"),
    ("Slitting Information:", "Slitting Information"),
    ("Packing Information:", "Packing Information"),
    ("Reporting Requirements/", "Reporting Requirements"),
    # Misspellings.
    ("litting Information", "Slitting Information"),
    ("Packing Informatio", "Packing Information"),
    ("Physical Attribues & Testing", "Physical Attributes & Testing"),
    ("Location", "Locations"),
    # A bilingual plant names the section, then translates it.
    ("Packing Information/ Information d’emballage", "Packing Information"),
    ("Process Routing/ Étape de production", "Process Routing"),
    # Sections that were unclassifiable until they were named.
    ("Extruder Distribution", "Extruder Distribution"),
    ("Blown Film – Blender Verification on KIEFEL Line", "Blown Film Blender Verification"),
    ("Blown Film – Blender Verification on ALPINE Line", "Blown Film Blender Verification"),
]

# Headings that must stay unclassified. Some are real sections of their own
# that nobody has asked Cobalt to understand; some only look close.
REFUSED = [
    "Quality Issues",
    "McGuires",
    "General Information",
    "Finish Good Information",
    "Finish Goods Information",
    "FG Section",
    "Packing Details",
    "Press Packing Information",
    "Materials",
    "Slitting(out-source ) Information",
]


@pytest.mark.parametrize("heading,expected", RECOVERED)
def test_a_heading_the_archive_really_uses_is_placed(heading, expected):
    section, _ = classify_heading(heading)
    assert section == expected


@pytest.mark.parametrize("heading", REFUSED)
def test_a_heading_cobalt_does_not_know_is_not_guessed_at(heading):
    section, _ = classify_heading(heading)
    assert section is None, f"{heading!r} was guessed as {section!r}"


def test_the_two_blown_film_lines_are_variants_of_one_section():
    """Same section, two machines -- the Duplex/Triplex shape, so both can
    live in one spec and stay separately addressable."""
    kiefel = classify_heading("Blown Film – Blender Verification on KIEFEL Line")
    alpine = classify_heading("Blown Film – Blender Verification on ALPINE Line")
    assert kiefel[0] == alpine[0] == "Blown Film Blender Verification"
    assert kiefel[1] == "KIEFEL Line"
    assert alpine[1] == "ALPINE Line"


def test_every_canonical_name_still_matches_itself():
    for name in BODY_SECTIONS:
        assert classify_heading(name) == (name, "")


def test_existing_variants_are_unaffected():
    assert classify_heading("Process Routing - Duplex") == ("Process Routing", "Duplex")
    assert classify_heading("Process Routing – Triplex") == ("Process Routing", "Triplex")


def test_the_old_aliases_still_work():
    assert classify_heading("Slitting Instructions")[0] == "Slitting Information"
    assert classify_heading("Packing Specifications")[0] == "Packing Information"


def test_two_sections_within_reach_are_refused_rather_than_picked():
    """The rule that keeps a near-miss from being resolved by luck."""
    from cobalt.docx_sections import _fuzzy_section

    assert _fuzzy_section("zzzzzzzzzzzzzzzzzzzzzzzz") is None


def test_a_human_override_still_wins_over_a_fuzzy_match():
    section, _ = classify_heading("Packing Informatio", overrides={"packing informatio": "Locations"})
    assert section == "Locations"


def test_the_new_sections_parse_out_of_a_real_document(tmp_path):
    """Classification is only half of it -- the tables have to come through
    as addressable rows, including both Blown Film lines in one spec."""
    from docx import Document

    from cobalt.docx_sections import parse_document

    path = str(tmp_path / "extruder.docx")
    doc = Document()
    doc.add_paragraph("Extruder Distribution")
    t = doc.add_table(rows=2, cols=3)
    for c, text in enumerate(["Extruder", "Resin", "Percent"]):
        t.cell(0, c).text = text
    for c, text in enumerate(["A", "LLDPE", "40"]):
        t.cell(1, c).text = text

    for line, resin in [("KIEFEL", "LDPE"), ("ALPINE", "HDPE")]:
        doc.add_paragraph(f"Blown Film – Blender Verification on {line} Line")
        t = doc.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "Blender"
        t.cell(0, 1).text = "Resin"
        t.cell(1, 0).text = line
        t.cell(1, 1).text = resin
    doc.save(path)

    spec = parse_document(path)
    assert "Extruder Distribution" in spec.tables
    assert spec.primary("Extruder Distribution").records()[0]["Resin"] == "LLDPE"

    blown = spec.tables["Blown Film Blender Verification"]
    assert [t.variant for t in blown] == ["KIEFEL Line", "ALPINE Line"]
    # Distinct table_index is what lets a write reach the right one.
    assert blown[0].table_index != blown[1].table_index
    assert [t.records()[0]["Resin"] for t in blown] == ["LDPE", "HDPE"]

    # Neither is reported as an unclassified table any more.
    assert [t.heading for t in spec.unclassified] == []
