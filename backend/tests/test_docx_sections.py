from cobalt.docx_sections import ALL_SECTIONS, OPTIONAL_SECTIONS, parse_document
from cobalt.models import TableShape

from .fixtures.builder import build_sample_spec_docx


def test_parses_all_sections(tmp_path):
    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path)

    spec = parse_document(path)

    assert spec.spec_number == "SW0001"
    assert spec.customer == "ACME Corp"
    assert spec.revision_number == "01"
    assert spec.warnings == []
    # Everything except the sections only some plants carry -- this fixture
    # is an ordinary spec, and not having an Extruder Distribution table is
    # the normal case rather than a parse failure.
    assert set(spec.tables) == set(ALL_SECTIONS) - OPTIONAL_SECTIONS


def test_records_shape_for_bom(tmp_path):
    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path)
    spec = parse_document(path)

    bom = spec.primary("Bill of Materials")
    assert bom.shape == TableShape.RECORDS
    assert bom.header_row == ["Caliper (mils)", "Raw Material", "Supplier", "Designation", "Part Number"]
    records = bom.records()
    assert len(records) == 1
    assert records[0]["Raw Material"] == "48g PET"
    assert records[0]["Supplier"] == "Flex Films"


def test_fields_shape_for_locations(tmp_path):
    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path)
    spec = parse_document(path)

    locations = spec.primary("Locations")
    assert locations.shape == TableShape.FIELDS
    fields = locations.fields()
    assert fields["Customer Location"] == "Test Customer"
    assert fields["Facility"] == "Test Plant"


def test_product_description_from_header_with_merged_cell(tmp_path):
    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path)
    spec = parse_document(path)

    pd = spec.primary("Product Description")
    assert pd.location == "header"
    fields = pd.fields()
    assert fields["Spec #"] == "SW0001"
    assert fields["Revision #"] == "01"
    assert fields["Structure Code"] == "TC-001"
    # merged "Structure Description" label collapses to one entry, not two
    assert fields["Structure Description"] == "48g PET / Ink / Adh / 2.0 mil PE"


def test_revision_history_records(tmp_path):
    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path)
    spec = parse_document(path)

    rev = spec.primary("Revision History")
    assert rev.shape == TableShape.RECORDS
    records = rev.records()
    assert records[0]["Revision #"] == "01"
    assert records[0]["Revision"] == "Spec created."


def _spec_with_pd_labels(path: str, spec_label: str, extra_rows=None):
    """Minimal doc whose Product Description header uses a given spec-number
    label, for exercising the label variants real archive specs use."""
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    rows = [["Customer:", "ACME Corp", spec_label, "HK0070"]]
    rows += extra_rows or []
    rows.append(["Item:", "Pouch", "Revision #:", "4"])
    table = doc.sections[0].header.add_table(rows=len(rows), cols=4, width=Inches(6))
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            table.cell(r, c).text = text
    doc.save(path)


def test_spec_number_from_pre_rebrand_sonoco_label(tmp_path):
    """Specs written before the Sonoco->Toppan rename say "Sonoco Spec #".
    Reading them as a blank spec number breaks the sidebar, the Mass Edit
    Spec Number column, and every audit-log entry for that file."""
    path = str(tmp_path / "sonoco.docx")
    _spec_with_pd_labels(path, "Sonoco Spec #:")
    assert parse_document(path).spec_number == "HK0070"


def test_spec_number_from_toppan_label(tmp_path):
    path = str(tmp_path / "toppan.docx")
    _spec_with_pd_labels(path, "Toppan Spec #:")
    assert parse_document(path).spec_number == "HK0070"


def test_customer_spec_number_is_never_mistaken_for_ours(tmp_path):
    """"Customer Spec #" is the customer's own number (often blank) and must
    not be picked up as the org's spec number."""
    path = str(tmp_path / "cust.docx")
    _spec_with_pd_labels(
        path, "Sonoco Spec #:", extra_rows=[["Customer Spec #:", "CUST-999", "", ""]]
    )
    assert parse_document(path).spec_number == "HK0070"


def test_section_title_aliases_are_recognized(tmp_path):
    """Pouch specs title these sections differently than roll specs. Without
    alias handling their tables are invisible to the app even though the
    data is right there in the document."""
    from docx import Document

    path = str(tmp_path / "aliases.docx")
    doc = Document()
    doc.add_paragraph("Slitting Instructions")
    t = doc.add_table(rows=1, cols=4)
    t.cell(0, 0).text = "Slit Width:"
    t.cell(0, 1).text = "20.25\""
    t.cell(0, 2).text = "Core ID:"
    t.cell(0, 3).text = "3\""
    doc.add_paragraph("Packing Specifications")
    t2 = doc.add_table(rows=2, cols=2)
    t2.cell(0, 0).text = "Pallet ID"
    t2.cell(0, 1).text = "Box Size"
    t2.cell(1, 0).text = "48 x 40"
    t2.cell(1, 1).text = "LH-2"
    doc.save(path)

    spec = parse_document(path)
    assert "Slitting Information" in spec.tables
    assert "Packing Information" in spec.tables
    assert spec.primary("Slitting Information").fields()["Slit Width"] == "20.25\""
