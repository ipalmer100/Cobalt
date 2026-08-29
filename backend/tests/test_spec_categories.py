"""Spec categories, and what each one is offered.

The extrusion plants' specs are a different kind of document, built around
sections no other spec has. They are their own category, and Mass Edit is
the standard category only -- a blown film spec is edited one at a time in
Spec Detail. These tests hold that line from both ends: the category is
read from the document rather than configured, and no blown film row
reaches a grid.
"""

from docx import Document

from cobalt.docx_sections import (
    CATEGORY_BLOWN_FILM,
    CATEGORY_STANDARD,
    STANDARD_SECTIONS,
    categorizing_sections,
    parse_document,
    spec_category,
)
from cobalt.vault import VaultEntry
from cobalt.views import VIEW_NAMES, build_view

from .fixtures.builder import build_sample_spec_docx

EXPECTED_VIEWS = [
    "Product Description",
    "Locations",
    "Bill of Materials",
    "Secondary Approved Materials",
    "Process Routing",
    "Physical Attributes & Testing",
    "Slitting Information",
    "Packing Information",
    "Reporting Requirements",
    "Customer Sampling Requirements",
    "Revision History",
]


def _blown_film(path, spec_number="BF0001"):
    """A standard spec plus the section that makes it blown film."""
    build_sample_spec_docx(str(path), spec_number=spec_number)
    doc = Document(str(path))
    doc.add_paragraph("Blown Film – Blender Verification on KIEFEL Line")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Blender"
    t.cell(0, 1).text = "Resin"
    t.cell(1, 0).text = "KIEFEL"
    t.cell(1, 1).text = "HDPE"
    doc.save(str(path))


def _entries(*paths):
    return [VaultEntry(path=str(p), spec=parse_document(str(p)), error=None, supported=True) for p in paths]


def test_mass_edit_offers_exactly_the_standard_sections():
    assert VIEW_NAMES == EXPECTED_VIEWS
    assert STANDARD_SECTIONS == EXPECTED_VIEWS


def test_no_category_specific_section_is_offered_as_a_view():
    assert "Blown Film Blender Verification" not in VIEW_NAMES
    assert "Extruder Distribution" not in VIEW_NAMES


def test_an_ordinary_spec_is_standard(tmp_path):
    path = tmp_path / "spec.docx"
    build_sample_spec_docx(str(path))
    spec = parse_document(str(path))
    assert spec_category(spec) == CATEGORY_STANDARD
    assert categorizing_sections(spec) == []


def test_a_spec_carrying_a_blown_film_section_is_blown_film(tmp_path):
    path = tmp_path / "bf.docx"
    _blown_film(path)
    spec = parse_document(str(path))
    assert spec_category(spec) == CATEGORY_BLOWN_FILM
    assert categorizing_sections(spec) == ["Blown Film Blender Verification"]


def test_an_extruder_distribution_spec_is_blown_film_too(tmp_path):
    path = tmp_path / "ex.docx"
    build_sample_spec_docx(str(path))
    doc = Document(str(path))
    doc.add_paragraph("Extruder Distribution")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Extruder"
    t.cell(0, 1).text = "Resin"
    t.cell(1, 0).text = "A"
    t.cell(1, 1).text = "LLDPE"
    doc.save(str(path))

    spec = parse_document(str(path))
    assert spec_category(spec) == CATEGORY_BLOWN_FILM
    assert categorizing_sections(spec) == ["Extruder Distribution"]


def test_a_blown_film_spec_contributes_no_rows_to_any_view(tmp_path):
    """The point of the category: nothing it holds can be reached by a
    fill, a bulk revision, or anything else the grid can do."""
    standard = tmp_path / "std.docx"
    blown = tmp_path / "bf.docx"
    build_sample_spec_docx(str(standard), spec_number="SW0001")
    _blown_film(blown, spec_number="BF0001")
    entries = _entries(standard, blown)

    for view in VIEW_NAMES:
        rows = build_view(entries, view)
        numbers = {row.get("Spec Number") for row in rows}
        assert "BF0001" not in numbers, f"{view} leaked a blown film row"

    # ...while the standard spec is untouched by any of this.
    assert {r["Spec Number"] for r in build_view(entries, "Bill of Materials")} == {"SW0001"}


def test_secondary_approved_materials_has_a_view_of_its_own(tmp_path):
    path = tmp_path / "spec.docx"
    build_sample_spec_docx(str(path))
    rows = build_view(_entries(path), "Secondary Approved Materials")
    assert rows, "Secondary Approved Materials should be a view in its own right"
    assert all(row["Spec Number"] == "SW0001" for row in rows)


def test_bill_of_materials_still_unions_both_lists(tmp_path):
    """Unchanged: the sample workbook and the VBA extractor both treat
    Primary and Secondary as one editable list."""
    path = tmp_path / "spec.docx"
    build_sample_spec_docx(str(path))
    rows = build_view(_entries(path), "Bill of Materials")
    assert {row["Material Type"] for row in rows} == {"Primary", "Secondary"}


def test_the_api_reports_a_category_for_every_spec(tmp_path):
    from fastapi.testclient import TestClient

    from cobalt.api import app

    standard = tmp_path / "std.docx"
    blown = tmp_path / "bf.docx"
    build_sample_spec_docx(str(standard), spec_number="SW0001")
    _blown_film(blown, spec_number="BF0001")

    client = TestClient(app)
    client.post("/vault/open", json={"root": str(tmp_path)})

    by_number = {e["spec_number"]: e for e in client.get("/vault").json()["entries"]}
    assert by_number["SW0001"]["category"] == CATEGORY_STANDARD
    assert by_number["BF0001"]["category"] == CATEGORY_BLOWN_FILM
    assert by_number["BF0001"]["category_sections"] == ["Blown Film Blender Verification"]

    detail = client.get("/spec", params={"path": str(blown)}).json()
    assert detail["category"] == CATEGORY_BLOWN_FILM
    assert detail["category_sections"] == ["Blown Film Blender Verification"]

    assert client.get("/views").json()["categories"] == {
        "standard": "Standard",
        "blown-film": "Blown Film",
    }


def test_a_blown_film_spec_is_still_fully_readable_and_editable_on_its_own(tmp_path):
    """It loses bulk treatment, not its data: Spec Detail must show every
    section, including the ones that put it in this category."""
    from fastapi.testclient import TestClient

    from cobalt.api import app

    blown = tmp_path / "bf.docx"
    _blown_film(blown)
    client = TestClient(app)
    client.post("/vault/open", json={"root": str(tmp_path)})

    detail = client.get("/spec", params={"path": str(blown)}).json()
    assert "Blown Film Blender Verification" in detail["sections"]
    assert "Bill of Materials" in detail["sections"]

    written = client.post(
        "/spec/commit",
        json={
            "who": "Isaac",
            "revision_text": "Blender resin corrected.",
            "edits": [
                {
                    "path": str(blown),
                    "section": "Blown Film Blender Verification",
                    "kind": "record",
                    "row": 1,
                    "col": 1,
                    "value": "LDPE",
                    "table_index": detail["sections"]["Blown Film Blender Verification"][0][
                        "table_index"
                    ],
                }
            ],
        },
    )
    assert written.status_code == 200, written.text
    after = parse_document(str(blown))
    assert after.primary("Blown Film Blender Verification").records()[0]["Resin"] == "LDPE"
