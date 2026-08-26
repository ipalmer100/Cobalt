"""Multiple tables per section (Franklin's two-process-path specs) and the
exception queue that catches everything the classifier won't guess at."""

from docx import Document

from specwrite.docx_sections import IGNORE, classify_heading, parse_document
from specwrite.section_mappings import load_mappings, save_mapping
from specwrite.views import build_view
from specwrite.vault import VaultEntry


def _doc_with(path: str, blocks: list[tuple[str, list[list[str]]]]):
    doc = Document()
    for heading, rows in blocks:
        doc.add_paragraph(heading)
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        for r, row in enumerate(rows):
            for c, text in enumerate(row):
                table.cell(r, c).text = text
    doc.save(path)


ROUTING_DUPLEX = [["Pass", "Process", "Machine"], ["1", "Print", "Rotomec"]]
ROUTING_TRIPLEX = [["Pass", "Process", "Machine"], ["1", "Laminate", "Triplex"]]


def test_classify_recognizes_qualified_variants():
    assert classify_heading("Process Routing - Duplex") == ("Process Routing", "Duplex")
    assert classify_heading("Physical Attributes & Testing - Triplex") == (
        "Physical Attributes & Testing",
        "Triplex",
    )
    assert classify_heading("Slitting Information - IMS Dairy Product") == (
        "Slitting Information",
        "IMS Dairy Product",
    )


def test_classify_escalates_rather_than_guessing():
    """Anything that isn't clearly one of the 11 must go to a human, not get
    filed somewhere plausible-looking."""
    for heading in ["Press Specification", "Quality Issues", "S3 Machine Conditions"]:
        assert classify_heading(heading) == (None, ""), heading


def test_spec_keeps_every_table_for_a_section(tmp_path):
    path = str(tmp_path / "two_paths.docx")
    _doc_with(path, [("Process Routing - Duplex", ROUTING_DUPLEX),
                     ("Process Routing - Triplex", ROUTING_TRIPLEX)])

    tables = parse_document(path).tables["Process Routing"]
    assert [t.variant for t in tables] == ["Duplex", "Triplex"]
    # distinct physical tables -- this is what lets a write target one of them
    assert tables[0].table_index != tables[1].table_index


def test_view_merges_variants_into_the_one_section(tmp_path):
    path = str(tmp_path / "two_paths.docx")
    _doc_with(path, [("Process Routing - Duplex", ROUTING_DUPLEX),
                     ("Process Routing - Triplex", ROUTING_TRIPLEX)])
    spec = parse_document(path)
    entry = VaultEntry(path=path, spec=spec, error=None, supported=True)

    rows = build_view([entry], "Process Routing")

    assert [r["Variant"] for r in rows] == ["Duplex", "Triplex"]
    assert [r["Process"] for r in rows] == ["Print", "Laminate"]
    # Same row number in each table -- only table_index tells them apart, so
    # without it an edit to one would land in the other.
    assert rows[0]["_source"]["row"] == rows[1]["_source"]["row"]
    assert rows[0]["_source"]["table_index"] != rows[1]["_source"]["table_index"]


def test_variant_column_absent_when_nothing_has_one(tmp_path):
    path = str(tmp_path / "single.docx")
    _doc_with(path, [("Process Routing", ROUTING_DUPLEX)])
    spec = parse_document(path)
    rows = build_view([VaultEntry(path=path, spec=spec, error=None, supported=True)], "Process Routing")
    assert "Variant" not in rows[0]


def test_unrecognized_table_goes_to_the_exception_queue(tmp_path):
    path = str(tmp_path / "odd.docx")
    _doc_with(path, [("Press Specification", [["Speed", "Units"], ["650", "ft/min"]])])

    spec = parse_document(path)

    assert [u.heading for u in spec.unclassified] == ["Press Specification"]
    assert spec.unclassified[0].row_count == 2
    assert spec.unclassified[0].preview[0] == ["Speed", "Units"]


def test_assigned_heading_lands_in_its_section(tmp_path):
    """A decision recorded in the queue reclassifies the table on reparse."""
    path = str(tmp_path / "odd.docx")
    _doc_with(path, [("Press Specification", [["Pass", "Process", "Machine"], ["1", "Print", "P15"]])])

    save_mapping(str(tmp_path), "Press Specification", "Process Routing", who="tester")
    overrides = {k: m.section for k, m in load_mappings(str(tmp_path)).items()}
    spec = parse_document(path, overrides)

    assert spec.unclassified == []
    assert [t.heading for t in spec.tables["Process Routing"]] == ["Press Specification"]


def test_ignored_heading_disappears_from_the_queue(tmp_path):
    path = str(tmp_path / "odd.docx")
    _doc_with(path, [("Quality Issues", [["Issue", "Action"], ["x", "y"]])])

    save_mapping(str(tmp_path), "Quality Issues", IGNORE, who="tester")
    overrides = {k: m.section for k, m in load_mappings(str(tmp_path)).items()}
    spec = parse_document(path, overrides)

    assert spec.unclassified == []
    assert "Quality Issues" not in spec.tables


def test_banner_row_above_the_header_is_not_mistaken_for_it(tmp_path):
    """FR0282's duplex Process Routing opens with a merged "Comments:" band.
    Reading that as the header leaves every column unmapped and the rows
    blank in the grid."""
    path = str(tmp_path / "banner.docx")
    # Shaped like the real one: the merged band only fills the cells it
    # spans, leaving the rest of the row empty.
    _doc_with(path, [("Process Routing", [
        ["Comments:", "Comments:", "", "", ""],
        ["Pass", "Process", "Machine", "Process Conditions", "Comments"],
        ["1", "Print", "Rotomec", "TL-014", "12 hr cure"],
    ])])

    table = parse_document(path).primary("Process Routing")
    assert table.header_index == 1
    assert table.header_row == ["Pass", "Process", "Machine", "Process Conditions", "Comments"]
    assert table.records() == [
        {
            "Pass": "1",
            "Process": "Print",
            "Machine": "Rotomec",
            "Process Conditions": "TL-014",
            "Comments": "12 hr cure",
        }
    ]


def test_duplicate_header_columns_stay_separately_addressable(tmp_path):
    """A merged header cell reads out as the same text in every column it
    spans. Keying rows on that text collapsed two columns into one: the
    grid showed the last column's value but a write landed in the first,
    so an edit appeared not to stick. Reported against a Physical
    Attributes table with a paired tolerance column."""
    path = str(tmp_path / "dupe.docx")
    _doc_with(path, [("Physical Attributes & Testing", [
        ["Attribute", "Tolerance", "Tolerance", "", "Frequency"],
        ["Repeat", "(-1/32)", "(+1/32)", "5.25", "Once"],
    ])])

    table = parse_document(path).primary("Physical Attributes & Testing")
    labels = table.column_labels()

    # every physical column keeps its own identity
    assert labels == ["Attribute", "Tolerance", "Tolerance (2)", "Column 4", "Frequency"]
    assert len(set(labels)) == len(labels)

    record = table.records()[0]
    assert record["Tolerance"] == "(-1/32)"
    assert record["Tolerance (2)"] == "(+1/32)"
    assert record["Column 4"] == "5.25"

    # ...and the index the grid derives from a label is the physical column
    for label, expected_col in [("Tolerance", 1), ("Tolerance (2)", 2), ("Column 4", 3)]:
        assert labels.index(label) == expected_col


def test_specs_spelling_a_column_differently_stay_separately_addressable(tmp_path):
    """Two plants write the same Bill of Materials column differently --
    Hazelton "Basis Wt Range", Franklin "Basis Wt range" -- and end their
    tables with different headings entirely ("Designation" vs "Raw Material
    Item Code").

    The grid merges the case variants for display and keeps the genuinely
    different headings apart. It can only do that if the view hands it, per
    row, the full positional header of *that row's own* table: a row keyed
    or addressed by some other spec's spelling is a row whose values can't
    be shown and whose edits can't be written.
    """
    hazelton = str(tmp_path / "hazelton.docx")
    franklin = str(tmp_path / "franklin.docx")
    _doc_with(hazelton, [("Bill of Materials", [
        ["Basis Wt", "Basis Wt Range", "Raw Material", "Designation"],
        ["9.60", "± 0.96", "70ga Matte OPP", "TD18-T 70G"],
    ])])
    _doc_with(franklin, [("Bill of Materials", [
        ["Basis Wt", "Basis Wt range", "Raw Material", "Raw Material Item Code"],
        ["8.10", "± 0.81", "48ga PET", "RM-4417"],
    ])])

    entries = [
        VaultEntry(path=p, spec=parse_document(p), error=None, supported=True)
        for p in (hazelton, franklin)
    ]
    rows = build_view(entries, "Bill of Materials")

    # Each row carries its own table's header, positionally -- that is the
    # write address, so it must never be borrowed from the other spec.
    assert rows[0]["_source"]["header_row"] == ["Basis Wt", "Basis Wt Range", "Raw Material", "Designation"]
    assert rows[1]["_source"]["header_row"] == ["Basis Wt", "Basis Wt range", "Raw Material", "Raw Material Item Code"]

    # ...and each row's values are keyed by that same spelling, so nothing
    # is stranded under a heading its own spec never used.
    assert rows[0]["Basis Wt Range"] == "± 0.96"
    assert rows[0]["Designation"] == "TD18-T 70G"
    assert rows[1]["Basis Wt range"] == "± 0.81"
    assert rows[1]["Raw Material Item Code"] == "RM-4417"

    # The position a write targets differs per spec even for the column the
    # grid shows as one: resolving it against the wrong header would put
    # Hazelton's edit in Franklin's column.
    for row in rows:
        header = row["_source"]["header_row"]
        label = next(h for h in header if h.lower() == "basis wt range")
        assert header.index(label) == 1
        assert row[label].startswith("±")


def test_edit_to_a_duplicate_named_column_sticks(tmp_path):
    """End to end: write through the column the grid would target, and read
    back the value the grid would display."""
    from specwrite.docx_writer import write_cell

    path = str(tmp_path / "dupe.docx")
    _doc_with(path, [("Physical Attributes & Testing", [
        ["Attribute", "Tolerance", "Tolerance", "Frequency"],
        ["Repeat", "(-1/32)", "(+1/32)", "Once"],
    ])])
    section = "Physical Attributes & Testing"
    labels = parse_document(path).primary(section).column_labels()

    write_cell(path, section, 1, labels.index("Tolerance"), "0.031")

    record = parse_document(path).primary(section).records()[0]
    assert record["Tolerance"] == "0.031"       # what was typed
    assert record["Tolerance (2)"] == "(+1/32)"  # its neighbour untouched
