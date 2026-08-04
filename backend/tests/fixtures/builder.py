"""Builds a small synthetic spec .docx for tests, structurally faithful to
the real Toppan spec template but with no real customer/business data.
Generated at test time rather than checked in as a binary fixture.
"""

from __future__ import annotations

from docx import Document
from docx.shared import Inches, RGBColor

RED = RGBColor(0xFF, 0x00, 0x00)


def _add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.color.rgb = RED
    run.bold = True


def _fill_table(table, rows: list[list[str]]) -> None:
    for r, row_values in enumerate(rows):
        for c, value in enumerate(row_values):
            table.cell(r, c).text = value


def build_sample_spec_docx(path: str, spec_number: str = "SW0001", revision: str = "01") -> None:
    doc = Document()

    header = doc.sections[0].header
    pd_table = header.add_table(rows=4, cols=4, width=Inches(6))
    _fill_table(
        pd_table,
        [
            ["Customer:", "ACME Corp", "Spec #:", spec_number],
            ["Item:", "12.00\" W.W. x 8.00\" C.O.", "Date of Issue:", "01/01/2024"],
            ["", "", "Revision #:", revision],
            ["", "", "Structure Code:", "TC-001"],
        ],
    )
    # Merge the two "Structure Description" cells vertically (as in real specs)
    # before filling their text, so the merge doesn't duplicate content.
    pd_table.cell(2, 0).merge(pd_table.cell(3, 0))
    pd_table.cell(2, 1).merge(pd_table.cell(3, 1))
    pd_table.cell(2, 0).text = "Structure Description:"
    pd_table.cell(2, 1).text = "48g PET / Ink / Adh / 2.0 mil PE"

    _add_heading(doc, "Locations")
    t = doc.add_table(rows=1, cols=4)
    _fill_table(t, [["Customer Location:", "Test Customer", "Facility:", "Test Plant"]])

    _add_heading(doc, "Bill of Materials")
    t = doc.add_table(rows=2, cols=5)
    _fill_table(
        t,
        [
            ["Caliper (mils)", "Raw Material", "Supplier", "Designation", "Part Number"],
            ["0.48", "48g PET", "Flex Films", "F-CHC-12", "PN-100"],
        ],
    )

    _add_heading(doc, "Secondary Approved Materials")
    t = doc.add_table(rows=2, cols=5)
    _fill_table(
        t,
        [
            ["Caliper (mils)", "Raw Material", "Supplier", "Designation", "Part Number"],
            ["0.48", "48g PET", "Terphane", "T-ALT-12", "PN-200"],
        ],
    )

    _add_heading(doc, "Process Routing")
    t = doc.add_table(rows=2, cols=5)
    _fill_table(
        t,
        [
            ["Pass", "Process", "Machine", "Process Conditions", "Comments"],
            ["1", "Print", "PRS FLX", "", "UW #1"],
        ],
    )

    _add_heading(doc, "Physical Attributes & Testing")
    t = doc.add_table(rows=2, cols=8)
    _fill_table(
        t,
        [
            ["Attribute", "COA", "Units", "LSL", "LCL", "Target", "UCL", "USL"],
            ["Caliper", "", "mils", "2.4", "2.5", "2.6", "2.7", "2.8"],
        ],
    )

    _add_heading(doc, "Slitting Information")
    t = doc.add_table(rows=1, cols=4)
    _fill_table(t, [["Slit Width:", "12.00\"", "Core ID:", "3\" Standard"]])

    _add_heading(doc, "Packing Information")
    t = doc.add_table(rows=1, cols=4)
    _fill_table(t, [["Pallet Size:", "48 x 40", "Poly Bags:", "Yes"]])

    _add_heading(doc, "Reporting Requirements")
    t = doc.add_table(rows=1, cols=4)
    _fill_table(t, [["COA Required?:", "No", "COA Recipient(s):", "quality@example.com"]])

    _add_heading(doc, "Customer Sampling Requirements")
    t = doc.add_table(rows=1, cols=4)
    _fill_table(t, [["Samples Required?:", "No", "Send to:", ""]])

    _add_heading(doc, "Revision History")
    t = doc.add_table(rows=2, cols=4)
    _fill_table(
        t,
        [
            ["Revision #", "Who", "Date", "Revision"],
            [revision, "Test User", "01/01/2024", "Spec created."],
        ],
    )

    doc.save(path)
