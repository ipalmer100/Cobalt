import time

import pytest
from docx import Document

from cobalt.docx_sections import parse_document
from cobalt.docx_writer import apply_revision, write_cell, write_edits_batch, write_record_cell

from .fixtures.builder import build_sample_spec_docx


def test_write_record_cell_preserves_other_cells(tmp_path):
    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path)

    spec = parse_document(path)
    bom = spec.primary("Bill of Materials")

    doc = Document(path)
    table = doc.tables[bom.table_index]
    write_record_cell(table, 1, 2, "New Supplier LLC")
    doc.save(path)

    spec2 = parse_document(path)
    records = spec2.primary("Bill of Materials").records()
    assert records[0]["Supplier"] == "New Supplier LLC"
    assert records[0]["Raw Material"] == "48g PET"  # untouched


def test_apply_revision_bumps_number_and_appends_row(tmp_path):
    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path, revision="06")

    new_rev = apply_revision(path, who="Isaac Palmer", revision_text="Updated supplier.")
    assert new_rev == "07"

    spec = parse_document(path)
    assert spec.revision_number == "07"

    rev_records = spec.primary("Revision History").records()
    assert len(rev_records) == 2
    assert rev_records[-1]["Revision #"] == "07"
    assert rev_records[-1]["Who"] == "Isaac Palmer"
    assert rev_records[-1]["Revision"] == "Updated supplier."


def test_apply_revision_zero_pads_width(tmp_path):
    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path, revision="09")

    new_rev = apply_revision(path, who="Tester", revision_text="Bump.")
    assert new_rev == "10"


def test_write_edits_batch_applies_multiple_record_edits_in_one_file(tmp_path):
    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path)

    write_edits_batch(
        [
            {"path": path, "section": "Bill of Materials", "kind": "record", "row": 1, "col": 2, "value": "Supplier A"},
            {"path": path, "section": "Secondary Approved Materials", "kind": "record", "row": 1, "col": 2, "value": "Supplier B"},
        ]
    )

    spec = parse_document(path)
    assert spec.primary("Bill of Materials").records()[0]["Supplier"] == "Supplier A"
    assert spec.primary("Secondary Approved Materials").records()[0]["Supplier"] == "Supplier B"


def test_write_edits_batch_applies_field_edits(tmp_path):
    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path)

    write_edits_batch(
        [{"path": path, "section": "Locations", "kind": "field", "label": "Facility", "value": "New Plant"}]
    )

    spec = parse_document(path)
    assert spec.primary("Locations").fields()["Facility"] == "New Plant"


def test_write_edits_batch_spans_multiple_files(tmp_path):
    path_a = str(tmp_path / "a.docx")
    path_b = str(tmp_path / "b.docx")
    build_sample_spec_docx(path_a, spec_number="SW0001")
    build_sample_spec_docx(path_b, spec_number="SW0002")

    write_edits_batch(
        [
            {"path": path_a, "section": "Bill of Materials", "kind": "record", "row": 1, "col": 2, "value": "X"},
            {"path": path_b, "section": "Bill of Materials", "kind": "record", "row": 1, "col": 2, "value": "Y"},
        ]
    )

    assert parse_document(path_a).primary("Bill of Materials").records()[0]["Supplier"] == "X"
    assert parse_document(path_b).primary("Bill of Materials").records()[0]["Supplier"] == "Y"


def test_write_edits_batch_is_faster_than_sequential_single_writes(tmp_path):
    """The whole point of batching: one open+save per file, not per cell."""
    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path)
    n = 20

    t0 = time.perf_counter()
    for i in range(n):
        write_cell(path, "Bill of Materials", 1, 2, f"Sequential {i}")
    sequential_time = time.perf_counter() - t0

    t0 = time.perf_counter()
    write_edits_batch(
        [{"path": path, "section": "Bill of Materials", "kind": "record", "row": 1, "col": 2, "value": f"Batched {i}"} for i in range(n)]
    )
    batch_time = time.perf_counter() - t0

    assert batch_time < sequential_time / 2


def test_write_cell_preserves_embedded_newlines(tmp_path):
    """~10% of cells in real specs are multi-line (a BOM "Raw Material"
    naming both sides of a film, a splice instruction, an address). Writing
    such a value must keep its line breaks -- collapsing them silently
    rewrites the meaning of the spec."""
    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path)

    multiline = "50g PET-Non Coated Side\n50g PET-PVDC Coated Side"
    write_cell(path, "Bill of Materials", 1, 1, multiline)

    value = parse_document(path).primary("Bill of Materials").rows[1][1]
    assert value == multiline
    assert "\n" in value


def test_concurrent_writes_to_one_spec_do_not_corrupt_it(tmp_path):
    """Two saves overlapping on the same file used to destroy it.

    Every write here is read-modify-write over the whole .docx, and
    ``doc.save(path)`` truncates the file and then streams a fresh zip into
    it. Two of those at once interleave, and the result is a spec Word --
    and python-docx -- can no longer open ("File name in directory
    'word/_rels/document.xml.rels' and header ... differ").

    Reachable in ordinary use, and much more so now that Spec Detail is
    editable: commit a cell, click straight into the next one and commit
    that, and the second save begins before the first has finished.
    """
    import threading
    import zipfile

    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path)

    errors: list[BaseException] = []

    def hammer(worker: int) -> None:
        try:
            for i in range(6):
                write_cell(path, "Bill of Materials", 1, 2, f"w{worker}-{i}")
        except BaseException as exc:  # noqa: BLE001 - reported, not swallowed
            errors.append(exc)

    threads = [threading.Thread(target=hammer, args=(w,)) for w in range(4)]
    for t in threads:
        t.start()
    for t in threads:
        t.join()

    assert errors == []
    # Still a valid zip, and still a readable spec.
    assert zipfile.ZipFile(path).testzip() is None
    spec = parse_document(path)
    assert spec.primary("Bill of Materials").rows[1][2].startswith("w")
    # Nothing else got mangled on the way through.
    assert spec.primary("Bill of Materials").header_row == (
        parse_document(path).primary("Bill of Materials").header_row
    )


def test_a_failed_save_leaves_the_original_intact(tmp_path):
    """The document is replaced by an atomic rename, so an interrupted save
    can't leave a half-written spec where the customer's document was, and
    can't leave a temp file behind for the vault indexer to trip over."""
    import zipfile

    from cobalt.docx_writer import _save_atomically

    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path)
    write_cell(path, "Bill of Materials", 1, 2, "Original Supplier")
    before = open(path, "rb").read()

    boom = RuntimeError("disk full")

    class HalfWritingDoc:
        """Saves some bytes, then dies -- what a full disk or a killed
        process actually looks like."""

        def save(self, target):
            with open(target, "wb") as fh:
                fh.write(b"PK\x03\x04 truncated")
            raise boom

    try:
        _save_atomically(HalfWritingDoc(), path)
    except RuntimeError as exc:
        assert exc is boom
    else:
        raise AssertionError("expected the save to fail")

    assert open(path, "rb").read() == before
    assert zipfile.ZipFile(path).testzip() is None
    assert parse_document(path).primary("Bill of Materials").rows[1][2] == "Original Supplier"
    assert list(tmp_path.glob(".*cobalt-tmp")) == []


def test_commit_writes_edits_and_the_revision_together(tmp_path):
    """Revisions are manual but regulatorily required, so the unit of work
    is "these edits plus this revision statement" -- a spec must never hold
    edited values with no Revision History row accounting for them."""
    from cobalt.docx_writer import commit_with_revision

    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path, revision="04")

    new_revs = commit_with_revision(
        {path: [
            {"section": "Bill of Materials", "kind": "record", "row": 1, "col": 2, "value": "New Supplier"},
        ]},
        who="Isaac",
        revision_text="Supplier change per customer request.",
    )

    assert new_revs == {path: "05"}
    spec = parse_document(path)
    assert spec.primary("Bill of Materials").rows[1][2] == "New Supplier"
    assert spec.revision_number == "05"
    last_revision = spec.primary("Revision History").rows[-1]
    assert last_revision[0] == "05"
    assert last_revision[1] == "Isaac"
    assert "Supplier change per customer request." in last_revision[3]


def test_a_failed_commit_writes_nothing_at_all(tmp_path):
    """The whole point of buffering: if the save is interrupted, the specs
    are exactly as they were -- no edited cells, and no revision claiming
    an edit that isn't there."""
    import zipfile

    from cobalt.docx_writer import commit_with_revision

    good = str(tmp_path / "good.docx")
    bad = str(tmp_path / "bad.docx")
    build_sample_spec_docx(good, revision="04")
    build_sample_spec_docx(bad, revision="04")
    before = {p: open(p, "rb").read() for p in (good, bad)}

    with pytest.raises(ValueError):
        commit_with_revision(
            {
                good: [{"section": "Bill of Materials", "kind": "record", "row": 1, "col": 2, "value": "Applied?"}],
                # A section this spec hasn't got: fails while staging, after
                # the first spec has already been prepared.
                bad: [{"section": "No Such Section", "kind": "record", "row": 1, "col": 1, "value": "x"}],
            },
            who="Isaac",
            revision_text="Should not land.",
        )

    for p in (good, bad):
        assert open(p, "rb").read() == before[p], f"{p} was modified"
        assert zipfile.ZipFile(p).testzip() is None
        spec = parse_document(p)
        assert spec.revision_number == "04"
        assert "Should not land." not in str(spec.tables)
    # No staging files left behind for the vault indexer to trip over.
    assert list(tmp_path.glob(".*cobalt-tmp")) == []


def test_commit_spanning_several_specs_revises_each_one(tmp_path):
    """A mass edit touches many specs; each gets its own revision number
    bumped and the same statement recorded against it."""
    from cobalt.docx_writer import commit_with_revision

    paths = []
    for i, rev in enumerate(["01", "07", "12"]):
        p = str(tmp_path / f"spec{i}.docx")
        build_sample_spec_docx(p, revision=rev)
        paths.append(p)

    new_revs = commit_with_revision(
        {p: [{"section": "Bill of Materials", "kind": "record", "row": 1, "col": 2, "value": "Acme"}] for p in paths},
        who="Isaac",
        revision_text="Rolled supplier across the family.",
    )

    assert [new_revs[p] for p in paths] == ["02", "08", "13"]
    for p in paths:
        spec = parse_document(p)
        assert spec.primary("Bill of Materials").rows[1][2] == "Acme"
        assert "Rolled supplier across the family." in spec.primary("Revision History").rows[-1][3]


def test_commit_leaves_untouched_specs_byte_identical(tmp_path):
    """A batch must not rewrite specs it wasn't asked to change."""
    from cobalt.docx_writer import commit_with_revision

    edited = str(tmp_path / "edited.docx")
    bystander = str(tmp_path / "bystander.docx")
    build_sample_spec_docx(edited)
    build_sample_spec_docx(bystander)
    before = open(bystander, "rb").read()

    commit_with_revision(
        {edited: [{"section": "Bill of Materials", "kind": "record", "row": 1, "col": 2, "value": "Changed"}]},
        who="Isaac",
        revision_text="One spec only.",
    )

    assert open(bystander, "rb").read() == before


def test_revision_number_survives_a_trailing_blank_history_row(tmp_path):
    """HK0070's Revision History ends in an entirely blank row -- somebody
    pressed Tab once too often. Reading that as the previous revision
    restarted a spec sitting at revision 4 back at "01", destroying the
    sequence the history exists to establish."""
    from docx import Document as Doc

    from cobalt.docx_writer import apply_revision

    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path, revision="4")

    # Give it the trailing blank row the real spec has.
    doc = Doc(path)
    rev_index = parse_document(path).primary("Revision History").table_index
    doc.tables[rev_index].add_row()
    doc.save(path)
    assert parse_document(path).primary("Revision History").rows[-1] == ["", "", "", ""]

    new_rev = apply_revision(path, "Isaac", "Next one along.")

    assert new_rev == "5", "must continue the sequence, not restart it"
    spec = parse_document(path)
    assert spec.revision_number == "5"
    rows = spec.primary("Revision History").rows
    # The blank row is filled rather than left as a gap mid-history.
    assert rows[-1][0] == "5"
    assert rows[-1][3] == "Next one along."
    assert not any(all(c == "" for c in r) for r in rows), "no blank row left in the trail"


def test_revision_number_continues_from_whichever_source_is_ahead(tmp_path):
    """If Product Description and the history disagree, one of them missed a
    revision -- continuing from the lower would reissue a used number."""
    from docx import Document as Doc

    from cobalt.docx_writer import apply_revision

    path = str(tmp_path / "spec.docx")
    build_sample_spec_docx(path, revision="09")

    # History left behind at 02 while the stated revision moved on to 09.
    doc = Doc(path)
    rev_index = parse_document(path).primary("Revision History").table_index
    doc.tables[rev_index].rows[1].cells[0].text = "02"
    doc.save(path)

    assert apply_revision(path, "Isaac", "Catching up.") == "10"
