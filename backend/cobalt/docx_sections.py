"""Read-side parser: locate each spec's 11 sections inside a .docx and
flatten them into structured tables.

Detection strategy (generalized from the org's existing BOM-extraction VBA):
for each known section name, find the paragraph whose text matches it, then
take the *next table that follows it* in document order. Product Description
is the exception — it lives in the page header as a label/value grid, with a
fallback to the first body table for older specs that never got a header
(mirrors the VBA's ``headerNum = 0`` branch).
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from .models import ParsedTable, Spec, TableShape, UnclassifiedTable

# Canonical section names, in the order they appear in a spec.
PRODUCT_DESCRIPTION = "Product Description"

BODY_SECTIONS = [
    "Locations",
    "Bill of Materials",
    "Secondary Approved Materials",
    "Process Routing",
    "Physical Attributes & Testing",
    "Slitting Information",
    "Packing Information",
    "Reporting Requirements",
    "Customer Sampling Requirements",
    "Revision History",
]

ALL_SECTIONS = [PRODUCT_DESCRIPTION, *BODY_SECTIONS]

_FIELD_GRID_MIN_COLON_FRACTION = 0.25


def _collapse(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _normalize(text: str) -> str:
    return _collapse(text).lower()


_NORMALIZED_BODY_SECTIONS = {_normalize(name): name for name in BODY_SECTIONS}

# Real specs in the archive title the same section differently depending on
# their vintage/product line -- a pouch spec's "Slitting Instructions" is the
# roll spec's "Slitting Information". Without these the section's table is
# invisible to the app even though it's right there in the document.
_SECTION_ALIASES = {
    "slitting instructions": "Slitting Information",
    "packing specifications": "Packing Information",
    "packaging information": "Packing Information",
}

for _alias, _canonical in _SECTION_ALIASES.items():
    _NORMALIZED_BODY_SECTIONS[_alias] = _canonical

# A heading may qualify a canonical section rather than rename it --
# "Process Routing - Duplex" is still Process Routing, just the duplex one.
# Only these separators count, so "Process Routing Sheet Numbering" (a
# different thing that merely starts with the same words) is not swallowed.
_VARIANT_SEPARATORS = ("-", "–", "—", ":", "(", "‒")

# Sentinel a human can assign in the exception queue to mean "this table
# isn't one of the 11 -- stop asking about it".
IGNORE = "__ignore__"


def classify_heading(
    text: str, overrides: dict[str, str] | None = None
) -> tuple[str | None, str]:
    """Map a heading onto one of the canonical sections.

    Returns ``(section, variant)``; ``section`` is None when the heading
    cannot be matched confidently, which routes the table to the exception
    queue instead of being guessed at. ``overrides`` carries decisions a
    human already made there (normalized heading -> section, or IGNORE).
    """
    collapsed = _collapse(text)
    norm = collapsed.lower()
    if not norm:
        return None, ""

    if overrides:
        assigned = overrides.get(norm)
        if assigned == IGNORE:
            return IGNORE, ""
        if assigned:
            # Keep any qualifier so an assigned "Press Routing - Duplex"
            # still shows up as the Duplex rows within its section.
            _, variant = _match_known(collapsed) or (None, "")
            return assigned, variant

    match = _match_known(collapsed)
    if match:
        return match
    return None, ""


def _match_known(collapsed: str) -> tuple[str, str] | None:
    norm = collapsed.lower()
    if norm in _NORMALIZED_BODY_SECTIONS:
        return _NORMALIZED_BODY_SECTIONS[norm], ""
    if norm == _normalize(PRODUCT_DESCRIPTION):
        return PRODUCT_DESCRIPTION, ""

    # Longest key first so "Physical Attributes & Testing" wins over any
    # shorter key that happens to prefix it.
    for key in sorted(_NORMALIZED_BODY_SECTIONS, key=len, reverse=True):
        if not norm.startswith(key) or len(norm) <= len(key):
            continue
        remainder = collapsed[len(key):]
        if remainder[0] not in _VARIANT_SEPARATORS and not remainder[0].isspace():
            continue
        variant = remainder.strip().lstrip("".join(_VARIANT_SEPARATORS)).strip()
        variant = variant.rstrip(")").strip()
        if not variant:
            continue
        return _NORMALIZED_BODY_SECTIONS[key], variant
    return None


def _table_to_rows(table: Table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    return rows


def _classify_shape(rows: list[list[str]]) -> TableShape:
    cells = [c for row in rows for c in row if c]
    if not cells:
        return TableShape.RECORDS
    colon_count = sum(1 for c in cells if c.rstrip().endswith(":"))
    fraction = colon_count / len(cells)
    return TableShape.FIELDS if fraction >= _FIELD_GRID_MIN_COLON_FRACTION else TableShape.RECORDS


def _build_parsed_table(
    section: str,
    table: Table,
    table_index: int,
    location: str,
    variant: str = "",
    heading: str = "",
) -> ParsedTable:
    rows = _table_to_rows(table)
    shape = _classify_shape(rows)
    header_index = _header_row_index(rows) if shape == TableShape.RECORDS else 0
    header_row = rows[header_index] if shape == TableShape.RECORDS and rows else None
    return ParsedTable(
        section=section,
        table_index=table_index,
        location=location,
        shape=shape,
        header_row=header_row,
        rows=rows,
        variant=variant,
        heading=heading,
        header_index=header_index,
    )


# A merged banner row spanning the full width ("Comments:" across FR0282's
# duplex Process Routing) repeats one value into every cell. Taking it as
# the header leaves the real column names unread and every row blank in the
# grid, so look one row further down when the first row looks like that.
_MIN_HEADER_DISTINCT_VALUES = 3


def _header_row_index(rows: list[list[str]]) -> int:
    if len(rows) < 2:
        return 0
    first_distinct = {c for c in rows[0] if c}
    if len(first_distinct) > 1:
        return 0
    second_distinct = {c for c in rows[1] if c}
    if len(second_distinct) >= _MIN_HEADER_DISTINCT_VALUES:
        return 1
    return 0


def _iter_body_blocks(doc: Document):
    """Yield ('p', Paragraph) / ('tbl', Table, table_index) in document order."""
    body = doc.element.body
    paragraphs = doc.paragraphs
    tables = doc.tables
    p_idx = 0
    t_idx = 0
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            yield ("p", paragraphs[p_idx])
            p_idx += 1
        elif tag == "tbl":
            yield ("tbl", tables[t_idx], t_idx)
            t_idx += 1


# A paragraph is only treated as a candidate table heading if it looks like
# one: short, and not a sentence. Body prose sitting above a table would
# otherwise flood the exception queue with things nobody needs to triage.
_MAX_HEADING_CHARS = 60


def _is_heading_candidate(text: str) -> bool:
    collapsed = _collapse(text)
    if not collapsed or len(collapsed) > _MAX_HEADING_CHARS:
        return False
    return not collapsed.endswith(".")


def find_body_section_tables(
    doc: Document, overrides: dict[str, str] | None = None
) -> tuple[dict[str, list[ParsedTable]], list[UnclassifiedTable]]:
    """Pair each section title paragraph with the table that follows it.

    A section maps to a *list* of tables: a spec written for two process
    paths carries e.g. both "Process Routing - Duplex" and "Process Routing
    - Triplex", and both belong under Process Routing. Headings that can't
    be matched confidently are returned separately for the exception queue
    rather than being guessed into a section.
    """
    found: dict[str, list[ParsedTable]] = {}
    unclassified: list[UnclassifiedTable] = []
    pending: tuple[str, str, str] | None = None  # (section|None-marker, variant, heading)

    for block in _iter_body_blocks(doc):
        if block[0] == "p":
            paragraph: Paragraph = block[1]
            raw = _collapse(paragraph.text)
            if not _is_heading_candidate(raw):
                continue
            section, variant = classify_heading(raw, overrides)
            pending = (section or "", variant, raw)
        else:
            _, table, table_index = block
            if pending is not None:
                section, variant, heading = pending
                if section == IGNORE:
                    pass
                elif section:
                    found.setdefault(section, []).append(
                        _build_parsed_table(section, table, table_index, "body", variant, heading)
                    )
                else:
                    rows = _table_to_rows(table)
                    shape = _classify_shape(rows)
                    unclassified.append(
                        UnclassifiedTable(
                            heading=heading,
                            table_index=table_index,
                            shape=shape,
                            header_row=rows[0] if shape == TableShape.RECORDS and rows else None,
                            row_count=len(rows),
                            preview=[r[:8] for r in rows[:4]],
                        )
                    )
            pending = None

    return found, unclassified


def extract_product_description(doc: Document) -> ParsedTable:
    """Product Description normally lives in the page header's first table;
    fall back to the first body table when a spec has no header table at all
    (older specs, per the VBA's ``headerNum = 0`` branch)."""
    header = doc.sections[0].header
    if header.tables:
        return _build_parsed_table(PRODUCT_DESCRIPTION, header.tables[0], 0, "header")
    if doc.tables:
        return _build_parsed_table(PRODUCT_DESCRIPTION, doc.tables[0], 0, "body")
    return ParsedTable(
        section=PRODUCT_DESCRIPTION,
        table_index=-1,
        location="missing",
        shape=TableShape.FIELDS,
        header_row=None,
        rows=[],
    )


def _spec_number_from_fields(pd_fields: dict[str, str]) -> str:
    """The org's own spec number, whatever the label calls it.

    Newer specs label it "Spec #"; ones written before the Sonoco->Toppan
    rename say "Sonoco Spec #", and a post-rename spec could equally say
    "Toppan Spec #". "Customer Spec #" is a *different* number (the
    customer's own, often blank) and must never be mistaken for ours.
    """
    for label in ("Spec #", "Sonoco Spec #", "Toppan Spec #"):
        value = pd_fields.get(label, "").strip()
        if value:
            return value
    for label, value in pd_fields.items():
        if label.strip().lower().endswith("spec #") and not label.strip().lower().startswith("customer"):
            if value.strip():
                return value.strip()
    return ""


def parse_bytes(key: str, data: bytes, overrides: dict[str, str] | None = None) -> Spec:
    """Parse a spec held in memory, labelled with the store key it came from.

    The storage-neutral entry point: a SharePoint document arrives as bytes
    over HTTPS and never exists as a file, and this is also what gets handed
    to worker processes during a parallel index (bytes and a string pickle
    cheaply; an open store connection does not).
    """
    return _parse(io.BytesIO(data), key, overrides)


def parse_document(path: str, overrides: dict[str, str] | None = None) -> Spec:
    """Parse a spec from a filesystem path."""
    return _parse(path, str(path), overrides)


def _parse(source, key: str, overrides: dict[str, str] | None = None) -> Spec:  # noqa: ANN001
    doc = Document(source)

    tables: dict[str, list[ParsedTable]] = {}
    warnings: list[str] = []

    tables[PRODUCT_DESCRIPTION] = [extract_product_description(doc)]
    body_tables, unclassified = find_body_section_tables(doc, overrides)
    for section, found in body_tables.items():
        tables.setdefault(section, []).extend(found)

    for name in ALL_SECTIONS:
        if name not in tables:
            warnings.append(f"Section not found: {name}")

    pd_fields = tables[PRODUCT_DESCRIPTION][0].fields()
    spec_number = _spec_number_from_fields(pd_fields)
    customer = pd_fields.get("Customer", "")
    revision_number = pd_fields.get("Revision #", "")

    return Spec(
        file_path=key,
        spec_number=spec_number,
        customer=customer,
        revision_number=revision_number,
        tables=tables,
        warnings=warnings,
        unclassified=unclassified,
    )
